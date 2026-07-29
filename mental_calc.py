"""
Copyright 2026 Martin Groß <martin@cavedev.de>

Permission is hereby granted, free of charge, to any person obtaining a copy of this software and associated
documentation files (the “Software”), to deal in the Software without restriction, including without limitation the
rights to use, copy, modify, merge, publish, distribute, sublicense, and/or sell copies of the Software, and to permit
persons to whom the Software is furnished to do so, subject to the following conditions:

The above copyright notice and this permission notice shall be included in all copies or substantial portions of
the Software.

THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED, INCLUDING BUT NOT LIMITED TO THE
WARRANTIES OF MERCHANTABILITY, FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE AUTHORS OR
COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR
OTHERWISE, ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE SOFTWARE.
"""

# Script to generate a specified number of paper sheets with calculation
# exercises for addition, subtraction, multiplication and division.
# The script generates a Word document with the exercises and the solutions.
# For this script to work, the python-docx library must be installed.

# Import the required libraries
import datetime
from enum import Enum
from docx import Document
from docx.shared import Pt
import random
import math


# Enumeration for the types of exercises
class ExerciseType(Enum):
    ADD = "Addition"
    SUB = "Subtraction"
    MUL = "Multiplication"
    DIV = "Division"


# Enumeration for the position of the answer field
class AnswerFieldPosition(Enum):
    LEFT = 1
    MIDDLE = 2
    RIGHT = 3


# Default minimum and maximum values for the exercises depending on the exercise type
DEFAULT_LIMITS_FOR_TYPE = {
    ExerciseType.ADD: (1, 20),
    ExerciseType.SUB: (1, 20),
    ExerciseType.MUL: (1, 10),
    ExerciseType.DIV: (1, 10)
}


# Function to generate the exercises as a list of strings
# parameters are the number of exercises and the maximum result value for
# the specified exercise type
def generate_exercises(
        num_exercises: int,
        limits_for_type: dict[ExerciseType, tuple[int, int]]) -> list[str]:
    # init the list of exercises
    exercises = []

    # get the maximum value for the exercises
    number_of_exercise_types = len(limits_for_type)
    exercise_types = list(limits_for_type.keys())

    # unicode character for the multiplication and division sign
    # mul_sign = u'\u00D7' # x sign
    mult_sign = u'\u22C5'  # dot sign
    # div_sign = u'\u00F7'  # division sign with bar
    div_sign = ":" # german division sign

    # underscore for the answer field
    answer_field_underscore = '___'

    # loop over the number of exercises
    for i in range(num_exercises):
        # get the type of the exercise and the maximum value for the type
        exercise_type = exercise_types[i % number_of_exercise_types]
        min_value = limits_for_type[exercise_type][0]
        max_value = limits_for_type[exercise_type][1]

        # init the string variables for the exercise
        left_field = ''
        middle_field = ''
        right_field = ''

        # init the numbers and the operator
        # a = 0
        # b = 0
        operator = ''

        # randomly determine where the answer field is placed, either left (first operator), middle (second operator)
        # or right (after the equals sign)
        answer_field_position = random.choice(
            [AnswerFieldPosition.LEFT, AnswerFieldPosition.MIDDLE, AnswerFieldPosition.RIGHT])

        # Switch to generate the numbers for the exercise
        match exercise_type:

            # Generate addition exercises
            case ExerciseType.ADD:
                operator = '+'
                a = random.randint(min_value, max_value - min_value)
                b = random.randint(min_value, max_value - a)

                # check for the position of the answer field and set the fields accordingly
                match answer_field_position:
                    case AnswerFieldPosition.LEFT:
                        left_field = answer_field_underscore
                        middle_field = f'{b}'
                        right_field = f'{a + b}'
                    case AnswerFieldPosition.MIDDLE:
                        left_field = f'{a}'
                        middle_field = answer_field_underscore
                        right_field = f'{a + b}'
                    case AnswerFieldPosition.RIGHT:
                        left_field = f'{a}'
                        middle_field = f'{b}'
                        right_field = answer_field_underscore

            # Generate subtraction exercises
            case ExerciseType.SUB:
                operator = '-'
                a_temp = random.randint(min_value, max_value - min_value)
                b_temp = random.randint(min_value, max_value - a_temp + min_value)
                # b_temp = 7
                a = max(a_temp, b_temp)
                b = min(a_temp, b_temp)

                # check for the position of the answer field and set the fields accordingly
                match answer_field_position:
                    case AnswerFieldPosition.LEFT:
                        left_field = answer_field_underscore
                        middle_field = f'{b}'
                        right_field = f'{a - b}'
                    case AnswerFieldPosition.MIDDLE:
                        left_field = f'{a}'
                        middle_field = answer_field_underscore
                        right_field = f'{a - b}'
                    case AnswerFieldPosition.RIGHT:
                        left_field = f'{a}'
                        middle_field = f'{b}'
                        right_field = answer_field_underscore

            # Generate multiplication exercises
            case ExerciseType.MUL:
                operator = mult_sign
                a = random.randint(1, max_value)
                b = random.randint(1, max_value)
                # a = random.randint(1, int(math.sqrt(max_value)))
                # b = random.randint(1, max_value // a)

                # check for the position of the answer field and set the fields accordingly
                match answer_field_position:
                    case AnswerFieldPosition.LEFT:
                        left_field = answer_field_underscore
                        middle_field = f'{b}'
                        right_field = f'{a * b}'
                    case AnswerFieldPosition.MIDDLE:
                        left_field = f'{a}'
                        middle_field = answer_field_underscore
                        right_field = f'{a * b}'
                    case AnswerFieldPosition.RIGHT:
                        left_field = f'{a}'
                        middle_field = f'{b}'
                        right_field = answer_field_underscore

            # Generate division exercises
            case ExerciseType.DIV:
                operator = div_sign
                a_temp = random.randint(1, max_value)
                b_temp = random.randint(1, max_value)
                a = a_temp * b_temp
                b = [a_temp, b_temp][random.randint(0, 1)]

                # Check for the position of the answer field and set the fields accordingly
                match answer_field_position:
                    case AnswerFieldPosition.LEFT:
                        left_field = answer_field_underscore
                        middle_field = f'{b}'
                        div = int(a / b)
                        right_field = f'{div}'
                    case AnswerFieldPosition.MIDDLE:
                        left_field = f'{a}'
                        middle_field = answer_field_underscore
                        div = int(a / b)
                        right_field = f'{div}'
                    case AnswerFieldPosition.RIGHT:
                        left_field = f'{a}'
                        middle_field = f'{b}'
                        right_field = answer_field_underscore

        exercises.append(f'{left_field} {operator} {middle_field} = {right_field}')
    return exercises


# Function to generate a Word document with the exercises
def generate_word_document(
        exercises: list[str],
        number_of_columns: int,
        filename: str):
    # Create a new Word document
    document = Document()

    # Check for multiple pages
    exercises_per_page = number_of_columns * 17  # assuming 17 rows per page
    num_pages = math.ceil(len(exercises) / exercises_per_page)
    for page in range(num_pages):
        if page > 0:
            document.add_page_break()

        # Add the title of the document
        document.add_heading('Name:\t\t\t\t\t\tDatum:', 0)

        # Add the exercises to the document as a table with the given number of columns
        table = document.add_table(rows=1, cols=number_of_columns)

        # Loop over the exercises and add them to the table
        for i, exercise in enumerate(exercises[page * exercises_per_page: (page + 1) * exercises_per_page]):
            # Add a new row if the number of columns is reached
            if i > 0 and i % number_of_columns == 0:
                table.add_row()
                # row_cells = table.add_row().cells

            # Add the exercise to the table
            row_number = i // number_of_columns
            row_cells = table.rows[row_number].cells
            row_cells[i % number_of_columns].text = exercise
            row_cells[i % number_of_columns].paragraphs[0].runs[0].font.size = Pt(18)

    # Save the document
    document.save(filename)


# The main function generates the exercises, and the Word document
# parameter is the number of pages to generate
def main(num_pages: int = 2):
    # print('Mental Calculation Exercises')
    exercises = generate_exercises(
        51 * num_pages,
        {
            ExerciseType.ADD: (1, 20),
            ExerciseType.SUB: (1, 20),
            ExerciseType.MUL: (1, 10),
            ExerciseType.DIV: (1, 10)
        })

    # Generate a file name with the current date and time
    filename_with_timestamp = \
        datetime.datetime.now().strftime('%Y-%m-%d_%H-%M-%S') + '_mental_calc.docx'

    # Generate the Word document with the exercises
    generate_word_document(
        exercises, number_of_columns=3, filename=filename_with_timestamp)


# Check if the script is run as the main program
if __name__ == '__main__':
    main(num_pages = 4)